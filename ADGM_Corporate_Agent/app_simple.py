"""
ADGM Corporate Agent - Simplified Application (Works without heavy AI dependencies)
AI-powered legal document review and compliance checking for Abu Dhabi Global Market
"""

import os
import sys
import json
import re
from typing import List, Dict, Tuple
import gradio as gr
import logging
from datetime import datetime

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class SimpleDocumentParser:
    """Simplified document parser without heavy dependencies"""
    
    def __init__(self):
        self.supported_formats = ['.docx']
    
    def parse_document(self, file_path: str) -> Dict:
        """Parse a DOCX document and extract information"""
        try:
            from docx import Document
            doc = Document(file_path)
            
            # Extract all text content
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            text_content = '\n'.join(full_text)
            
            # Extract document information
            doc_info = {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'text_content': text_content,
                'word_count': len(text_content.split()),
                'paragraph_count': len(full_text),
                'table_count': len(doc.tables),
                'document_type': self._identify_document_type(text_content)
            }
            
            logger.info(f"Successfully parsed: {doc_info['file_name']}")
            return doc_info
            
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {e}")
            raise Exception(f"Failed to parse document: {str(e)}")
    
    def _identify_document_type(self, text_content: str) -> str:
        """Identify document type based on content"""
        text_lower = text_content.lower()
        
        # Document identification patterns
        patterns = {
            'articles_of_association': [
                'articles of association',
                'company regulations', 
                'share capital',
                'director powers'
            ],
            'board_resolution': [
                'board resolution',
                'resolved that',
                'directors resolve',
                'board of directors'
            ],
            'ubo_declaration': [
                'ultimate beneficial owner',
                'ubo declaration', 
                'beneficial ownership',
                '25%'
            ],
            'employment_contract': [
                'employment contract',
                'terms of employment',
                'employee',
                'employer',
                'salary'
            ],
            'incorporation_form': [
                'incorporation',
                'application for registration',
                'company registration'
            ]
        }
        
        # Score each type
        scores = {}
        for doc_type, keywords in patterns.items():
            score = sum(1 for keyword in keywords if keyword in text_lower)
            if score > 0:
                scores[doc_type] = score
        
        # Return highest scoring type
        if scores:
            return max(scores, key=scores.get)
        return 'unknown'

class SimpleValidator:
    """Simplified document validator"""
    
    def __init__(self):
        self.red_flag_patterns = {
            "jurisdiction_issues": {
                "patterns": [
                    r"uae federal court",
                    r"dubai court",
                    r"abu dhabi court(?!.*global market)",
                    r"federal law",
                    r"emirates law"
                ],
                "severity": "high",
                "message": "Incorrect jurisdiction - should reference ADGM Courts"
            },
            "missing_adgm_reference": {
                "patterns": [
                    r"(?i)governing law(?!.*adgm)(?!.*abu dhabi global market)"
                ],
                "severity": "high",
                "message": "Missing ADGM governing law reference"
            },
            "incomplete_signatures": {
                "patterns": [
                    r"signature.*blank",
                    r"\[signature\]",
                    r"sign here",
                    r"_+.*signature"
                ],
                "severity": "medium", 
                "message": "Incomplete signature sections found"
            }
        }
    
    def validate_document(self, doc_info: Dict) -> Dict:
        """Validate document and detect red flags"""
        text_content = doc_info.get('text_content', '')
        text_lower = text_content.lower()
        
        red_flags = []
        
        # Check for red flags
        for flag_type, config in self.red_flag_patterns.items():
            for pattern in config['patterns']:
                if re.search(pattern, text_lower):
                    red_flags.append({
                        'type': flag_type,
                        'severity': config['severity'],
                        'message': config['message'],
                        'suggestion': 'Review and correct this issue'
                    })
        
        # Check for ADGM reference
        if 'adgm' not in text_lower and 'abu dhabi global market' not in text_lower:
            red_flags.append({
                'type': 'missing_adgm_reference',
                'severity': 'high',
                'message': 'Document does not reference ADGM or Abu Dhabi Global Market',
                'suggestion': 'Add proper ADGM governing law clause'
            })
        
        return {
            'document_name': doc_info.get('file_name', 'Unknown'),
            'document_type': doc_info.get('document_type', 'unknown'),
            'red_flags': red_flags,
            'overall_status': 'requires_attention' if red_flags else 'compliant'
        }

class SimpleComplianceChecker:
    """Simplified compliance checker"""
    
    def __init__(self):
        self.process_requirements = {
            "company_incorporation": {
                "name": "Company Incorporation",
                "required_docs": [
                    "articles_of_association",
                    "board_resolution",
                    "incorporation_form", 
                    "ubo_declaration",
                    "register_members_directors"
                ]
            }
        }
    
    def generate_compliance_report(self, documents: List[Dict], validations: List[Dict]) -> Dict:
        """Generate compliance report"""
        doc_types = {doc.get('document_type', 'unknown') for doc in documents}
        doc_types.discard('unknown')
        
        # Identify process (simplified)
        if any(dtype in ['articles_of_association', 'board_resolution', 'ubo_declaration'] for dtype in doc_types):
            process_type = 'company_incorporation'
            process_name = 'Company Incorporation'
            required_docs = self.process_requirements['company_incorporation']['required_docs']
        else:
            process_type = 'general_review'
            process_name = 'General Document Review'
            required_docs = []
        
        # Check completeness
        if required_docs:
            present_docs = [doc for doc in required_docs if doc in doc_types]
            missing_docs = [doc for doc in required_docs if doc not in doc_types]
            completion_pct = len(present_docs) / len(required_docs) * 100
        else:
            present_docs = list(doc_types)
            missing_docs = []
            completion_pct = 100.0
        
        # Count issues
        total_red_flags = sum(len(v.get('red_flags', [])) for v in validations)
        high_issues = sum(1 for v in validations for rf in v.get('red_flags', []) if rf.get('severity') == 'high')
        
        # Determine status
        if completion_pct < 60:
            status = 'incomplete'
        elif high_issues > 2:
            status = 'major_issues'
        elif high_issues > 0 or missing_docs:
            status = 'requires_attention'
        else:
            status = 'ready_for_submission'
        
        return {
            'process_type': process_type,
            'process_name': process_name,
            'overall_status': status,
            'completeness': {
                'documents_uploaded': len(documents),
                'required_documents': len(required_docs),
                'present_documents': present_docs,
                'missing_documents': missing_docs,
                'completion_percentage': completion_pct
            },
            'issues_summary': {
                'total_red_flags': total_red_flags,
                'high_severity': high_issues
            },
            'recommendations': [
                'Review flagged issues and make necessary corrections',
                'Ensure all required documents are provided',
                'Verify ADGM compliance before submission'
            ]
        }

class SimpleRAGSystem:
    """Simplified Q&A system with built-in ADGM knowledge"""
    
    def __init__(self):
        self.knowledge = {
            'incorporation': """ADGM company incorporation requires:
            1. Articles of Association - company structure and governance
            2. Board Resolution - shareholders' decisions  
            3. Incorporation Application Form - official registration
            4. UBO Declaration - beneficial owners with 25%+ ownership
            5. Register of Members and Directors - company officers listing
            
            All documents must reference ADGM jurisdiction and law.""",
            
            'jurisdiction': """ADGM documents must specify ADGM jurisdiction:
            
            ✅ CORRECT: "Governed by ADGM law and subject to ADGM Courts"
            ❌ INCORRECT: References to UAE Federal Courts or Dubai Courts
            
            ADGM has separate legal framework from UAE federal law.""",
            
            'ubo': """UBO declarations must include:
            - Individuals with 25%+ ownership or control
            - Complete personal details and addresses
            - Passport copies for all nationalities
            - Signed accuracy confirmations""",
            
            'employment': """ADGM employment contracts must include:
            - Job title, duties, reporting structure
            - Salary, benefits, payment terms
            - Working hours (max 48 hours/week)  
            - Annual leave (min 30 days)
            - Probation period (max 6 months)
            - ADGM-compliant termination procedures"""
        }
    
    def generate_response(self, query: str, context: str = "") -> str:
        """Generate response based on query"""
        query_lower = query.lower()
        
        # Find relevant knowledge
        response_parts = []
        for topic, content in self.knowledge.items():
            if any(keyword in query_lower for keyword in topic.split()):
                response_parts.append(f"**{topic.title()} Information:**\n\n{content}\n\n")
        
        # Check for specific keywords
        if any(word in query_lower for word in ['incorporation', 'incorporate', 'company']):
            if 'incorporation' not in [p for p in response_parts]:
                response_parts.append(self.knowledge['incorporation'])
        
        if any(word in query_lower for word in ['jurisdiction', 'court', 'law']):
            if 'jurisdiction' not in [p for p in response_parts]:
                response_parts.append(self.knowledge['jurisdiction'])
        
        if any(word in query_lower for word in ['ubo', 'beneficial', 'owner']):
            if 'ubo' not in [p for p in response_parts]:
                response_parts.append(self.knowledge['ubo'])
        
        if any(word in query_lower for word in ['employment', 'contract', 'employee']):
            if 'employment' not in [p for p in response_parts]:
                response_parts.append(self.knowledge['employment'])
        
        if response_parts:
            response = "Based on ADGM regulations:\n\n" + "\n".join(response_parts)
            response += "\n💡 **Note**: This is guidance only. Consult legal professionals for official advice."
        else:
            response = f"""I understand you're asking about: "{query}"

Here are key ADGM requirements:

🏛️ **Jurisdiction**: Documents must reference ADGM Courts and ADGM law
📋 **Incorporation**: Requires Articles, Board Resolution, UBO Declaration, and other documents  
👥 **UBO**: Declare individuals with 25%+ ownership
💼 **Employment**: Use ADGM-compliant contract terms

For specific guidance, please consult ADGM official resources or legal professionals."""
        
        return response

class SimpleReportGenerator:
    """Simplified report generator"""
    
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def generate_json_report(self, analysis_results: Dict) -> str:
        """Generate JSON report"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"adgm_analysis_{timestamp}.json"
        filepath = os.path.join(self.output_dir, filename)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(analysis_results, f, indent=2, ensure_ascii=False)
        
        return filepath
    
    def generate_text_report(self, analysis_results: Dict) -> str:
        """Generate text report"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"adgm_report_{timestamp}.txt"
        filepath = os.path.join(self.output_dir, filename)
        
        compliance_report = analysis_results.get("compliance_report", {})
        
        content = f"""
ADGM CORPORATE AGENT - DOCUMENT ANALYSIS REPORT
==============================================

Generated: {datetime.now().strftime("%B %d, %Y at %H:%M:%S")}

EXECUTIVE SUMMARY
-----------------
Process: {compliance_report.get('process_name', 'Unknown')}
Status: {compliance_report.get('overall_status', 'Unknown').replace('_', ' ').upper()}
Documents: {analysis_results.get('documents_processed', 0)}
Completion: {compliance_report.get('completeness', {}).get('completion_percentage', 0):.1f}%

FINDINGS
--------
Red Flags: {compliance_report.get('issues_summary', {}).get('total_red_flags', 0)}
High Severity: {compliance_report.get('issues_summary', {}).get('high_severity', 0)}

RECOMMENDATIONS
---------------
{chr(10).join(f"• {rec}" for rec in compliance_report.get('recommendations', []))}

DISCLAIMER
----------
This report is for guidance only and does not constitute legal advice.
Consult qualified legal professionals for official ADGM submissions.
"""
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return filepath

class ADGMCorporateAgent:
    """Main ADGM Corporate Agent Application - Simplified Version"""
    
    def __init__(self):
        """Initialize the ADGM Corporate Agent"""
        self.output_dir = "output"
        os.makedirs(self.output_dir, exist_ok=True)
        
        # Initialize simplified components
        self.document_parser = SimpleDocumentParser()
        self.rag_system = SimpleRAGSystem()
        self.document_validator = SimpleValidator()
        self.compliance_checker = SimpleComplianceChecker()
        self.report_generator = SimpleReportGenerator(self.output_dir)
        
        logger.info("ADGM Corporate Agent (Simplified) initialized successfully")
    
    def process_documents(self, files, progress=gr.Progress()):
        """Process uploaded documents and generate analysis"""
        if not files:
            return "❌ No files uploaded. Please upload DOCX documents.", "", ""
        
        try:
            progress(0.1, desc="Starting document analysis...")
            
            uploaded_documents = []
            document_analyses = []
            
            # Process each file
            for i, file in enumerate(files):
                progress(0.2 + (0.5 * i / len(files)), desc=f"Processing {os.path.basename(file.name)}...")
                
                try:
                    # Parse document
                    doc_info = self.document_parser.parse_document(file.name)
                    uploaded_documents.append(doc_info)
                    
                    # Validate document
                    validation = self.document_validator.validate_document(doc_info)
                    document_analyses.append(validation)
                    
                except Exception as e:
                    logger.error(f"Error processing {file.name}: {e}")
                    continue
            
            progress(0.7, desc="Generating compliance report...")
            
            # Generate compliance report
            compliance_report = self.compliance_checker.generate_compliance_report(
                uploaded_documents, document_analyses
            )
            
            progress(0.9, desc="Creating reports...")
            
            # Create analysis results
            analysis_results = {
                "timestamp": datetime.now().isoformat(),
                "documents_processed": len(uploaded_documents),
                "compliance_report": compliance_report,
                "document_analyses": document_analyses
            }
            
            # Generate output files
            json_path = self.report_generator.generate_json_report(analysis_results)
            text_path = self.report_generator.generate_text_report(analysis_results)
            
            progress(1.0, desc="Analysis complete!")
            
            # Generate status message
            status = self._generate_status_message(compliance_report, len(uploaded_documents))
            
            return status, json_path, text_path
            
        except Exception as e:
            logger.error(f"Error in document processing: {e}")
            return f"❌ Error: {str(e)}", "", ""
    
    def answer_question(self, question: str, context: str = ""):
        """Answer ADGM-related questions"""
        if not question.strip():
            return "Please enter a question about ADGM requirements."
        
        try:
            response = self.rag_system.generate_response(question, context)
            return response
        except Exception as e:
            return f"Error: {str(e)}"
    
    def _generate_status_message(self, compliance_report: Dict, doc_count: int) -> str:
        """Generate user-friendly status message"""
        process_name = compliance_report.get("process_name", "Document Analysis")
        status = compliance_report.get("overall_status", "completed")
        completion = compliance_report.get("completeness", {}).get("completion_percentage", 0)
        missing = compliance_report.get("completeness", {}).get("missing_documents", [])
        
        status_icons = {
            "ready_for_submission": "✅",
            "requires_attention": "⚠️", 
            "major_issues": "❌",
            "incomplete": "📋",
            "completed": "📄"
        }
        
        icon = status_icons.get(status, "📄")
        
        message = f"""
{icon} **Analysis Complete**

📊 **Process**: {process_name}
📈 **Completion**: {completion:.1f}%
📑 **Documents Analyzed**: {doc_count}
"""
        
        if missing:
            message += f"""
📋 **Missing Documents**:
{chr(10).join(f"   • {doc.replace('_', ' ').title()}" for doc in missing[:3])}
"""
        
        issues = compliance_report.get("issues_summary", {})
        high_issues = issues.get("high_severity", 0)
        if high_issues > 0:
            message += f"""
⚠️ **High Priority Issues**: {high_issues}
"""
        
        return message

def create_interface():
    """Create the Gradio interface"""
    agent = ADGMCorporateAgent()
    
    with gr.Blocks(
        title="ADGM Corporate Agent",
        theme=gr.themes.Soft(),
        css="""
        .main-header { 
            background: linear-gradient(90deg, #1e40af, #3b82f6);
            color: white;
            padding: 2rem;
            border-radius: 10px;
            text-align: center;
            margin-bottom: 2rem;
        }
        """
    ) as demo:
        
        # Header
        gr.HTML("""
        <div class="main-header">
            <h1>🏛️ ADGM Corporate Agent</h1>
            <p>AI-Powered Legal Document Intelligence for Abu Dhabi Global Market</p>
            <p><em>Simplified version - works without heavy AI dependencies</em></p>
        </div>
        """)
        
        with gr.Tabs():
            # Document Analysis Tab
            with gr.TabItem("📄 Document Analysis"):
                gr.Markdown("""
                ### Upload ADGM Documents for Analysis
                
                **The system will:**
                - 🔍 Parse and analyze document content
                - 🚩 Detect red flags and compliance issues
                - ✅ Validate against ADGM requirements  
                - 📊 Generate comprehensive reports
                """)
                
                with gr.Row():
                    with gr.Column(scale=2):
                        files = gr.File(
                            label="📤 Upload DOCX Documents",
                            file_count="multiple",
                            file_types=[".docx"],
                            height=150
                        )
                        
                        analyze_btn = gr.Button(
                            "🔍 Analyze Documents",
                            variant="primary",
                            size="lg"
                        )
                    
                    with gr.Column(scale=1):
                        gr.Markdown("""
                        **Supported Documents:**
                        
                        📋 **Company Formation**
                        • Articles of Association
                        • Board Resolutions
                        • UBO Declarations
                        • Incorporation Forms
                        
                        💼 **Other Documents**
                        • Employment Contracts
                        • Commercial Agreements
                        • Compliance Policies
                        """)
                
                # Results
                status_output = gr.Textbox(
                    label="📊 Analysis Results",
                    lines=8,
                    interactive=False
                )
                
                with gr.Row():
                    json_report = gr.File(label="📋 JSON Report")
                    text_report = gr.File(label="📄 Text Report")
                
                analyze_btn.click(
                    agent.process_documents,
                    inputs=[files],
                    outputs=[status_output, json_report, text_report]
                )
            
            # Q&A Assistant Tab
            with gr.TabItem("💬 ADGM Assistant"):
                gr.Markdown("""
                ### Ask Questions About ADGM Requirements
                
                Get expert guidance on ADGM regulations, processes, and compliance requirements.
                """)
                
                with gr.Row():
                    with gr.Column(scale=2):
                        question = gr.Textbox(
                            label="❓ Your Question",
                            placeholder="e.g., What documents are required for ADGM company incorporation?",
                            lines=3
                        )
                        
                        ask_btn = gr.Button("🤔 Get Answer", variant="primary")
                    
                    with gr.Column(scale=1):
                        gr.Markdown("""
                        **Example Questions:**
                        - Incorporation requirements
                        - UBO declaration rules  
                        - Jurisdiction requirements
                        - Employment contract terms
                        """)
                
                answer = gr.Textbox(
                    label="💡 Expert Answer",
                    lines=12,
                    interactive=False
                )
                
                ask_btn.click(
                    agent.answer_question,
                    inputs=[question],
                    outputs=[answer]
                )
                
                # Quick question buttons
                with gr.Row():
                    gr.Button("📋 Incorporation Docs").click(
                        lambda: agent.answer_question("What documents are required for company incorporation?"),
                        outputs=[answer]
                    )
                    
                    gr.Button("👥 UBO Rules").click(
                        lambda: agent.answer_question("What are UBO declaration requirements?"),
                        outputs=[answer]
                    )
                    
                    gr.Button("⚖️ Jurisdiction").click(
                        lambda: agent.answer_question("What jurisdiction should ADGM documents reference?"),
                        outputs=[answer]
                    )
            
            # About Tab
            with gr.TabItem("ℹ️ About"):
                gr.Markdown("""
                ## ADGM Corporate Agent - Simplified Version
                
                This is a simplified version that works without heavy AI dependencies like PyTorch and sentence-transformers.
                
                ### 🎯 Features
                
                - **Document Analysis**: Parse DOCX files and identify document types
                - **Red Flag Detection**: Find jurisdiction issues, missing ADGM references
                - **Compliance Checking**: Validate document completeness
                - **Q&A System**: Get answers about ADGM requirements
                - **Report Generation**: Create JSON and text reports
                
                ### 🚀 Quick Start
                
                1. Upload DOCX documents using the Document Analysis tab
                2. Review the analysis results and reports
                3. Ask questions using the ADGM Assistant
                
                ### ⚠️ Important Notes
                
                - This version has simplified AI capabilities
                - For production use, consider the full version with OpenAI integration
                - Always consult legal professionals for official guidance
                
                ### 📁 Files Generated
                
                Reports are saved in the `output/` directory:
                - JSON reports with structured data
                - Text reports with formatted summaries
                
                ---
                *Ready to analyze your ADGM documents!*
                """)
    
    return demo

if __name__ == "__main__":
    # Create and launch the application
    interface = create_interface()
    interface.launch(
        server_name="0.0.0.0",
        server_port=7860,
        debug=True,
        share=False
    )