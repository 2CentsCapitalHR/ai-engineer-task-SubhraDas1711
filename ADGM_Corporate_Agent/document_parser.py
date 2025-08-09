"""
Document Parser for ADGM Corporate Agent
Handles DOCX document parsing and content extraction
"""

import os
import re
from typing import Dict, List, Optional
from docx import Document
import logging

logger = logging.getLogger(__name__)

class DocumentParser:
    """DOCX document parser and analyzer"""
    
    def __init__(self):
        self.supported_formats = ['.docx']
    
    def parse_document(self, file_path: str) -> Dict:
        """Parse a DOCX document and extract information"""
        try:
            doc = Document(file_path)
            
            # Extract all text content
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            text_content = '\n'.join(full_text)
            
            # Extract document information
            doc_info = {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'text_content': text_content,
                'word_count': len(text_content.split()),
                'paragraph_count': len(full_text),
                'table_count': len(doc.tables),
                'document_type': self._identify_document_type(text_content),
                'metadata': self._extract_metadata(doc)
            }
            
            logger.info(f"Successfully parsed: {doc_info['file_name']}")
            return doc_info
            
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {e}")
            raise Exception(f"Failed to parse document: {str(e)}")
    
    def _identify_document_type(self, text_content: str) -> str:
        """Identify document type based on content"""
        text_lower = text_content.lower()
        
        # Document identification patterns
        patterns = {
            'articles_of_association': [
                'articles of association',
                'company regulations', 
                'share capital',
                'director powers'
            ],
            'board_resolution': [
                'board resolution',
                'resolved that',
                'directors resolve',
                'board of directors'
            ],
            'ubo_declaration': [
                'ultimate beneficial owner',
                'ubo declaration', 
                'beneficial ownership',
                '25%'
            ],
            'employment_contract': [
                'employment contract',
                'terms of employment',
                'employee',
                'employer',
                'salary'
            ],
            'incorporation_form': [
                'incorporation',
                'application for registration',
                'company registration'
            ]
        }
        
        # Score each type
        scores = {}
        for doc_type, keywords in patterns.items():
            score = sum(1 for keyword in keywords if keyword in text_lower)
            if score > 0:
                scores[doc_type] = score
        
        # Return highest scoring type
        if scores:
            return max(scores, key=scores.get)
        return 'unknown'
    
    def _extract_metadata(self, doc: Document) -> Dict:
        """Extract document metadata"""
        try:
            props = doc.core_properties
            return {
                'title': props.title or '',
                'author': props.author or '',
                'created': props.created.isoformat() if props.created else '',
                'modified': props.modified.isoformat() if props.modified else ''
            }
        except Exception as e:
            logger.warning(f"Could not extract metadata: {e}")
            return {}
    
    def validate_file(self, file_path: str) -> tuple[bool, str]:
        """Validate if file can be processed"""
        if not os.path.exists(file_path):
            return False, "File does not exist"
        
        _, ext = os.path.splitext(file_path.lower())
        if ext not in self.supported_formats:
            return False, f"Unsupported format. Only DOCX files are supported."
        
        try:
            Document(file_path)
            return True, "Valid DOCX file"
        except Exception as e:
            return False, f"Corrupted or invalid DOCX file: {str(e)}"